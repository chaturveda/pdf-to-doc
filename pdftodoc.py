import tkinter as tk
from tkinter import filedialog
from pdf2image import convert_from_path
import PyPDF2
from docx import Document
from docx.shared import Pt
import pytesseract

def select_pdf_file():
    file_path = filedialog.askopenfilename(filetypes=[("PDF Files", "*.pdf")])
    pdf_entry.delete(0, tk.END)
    pdf_entry.insert(0, file_path)

def select_output_folder():
    folder_path = filedialog.askdirectory()
    output_entry.delete(0, tk.END)
    output_entry.insert(0, folder_path)

def select_save_location():
    save_folder = filedialog.askdirectory()
    save_entry.delete(0, tk.END)
    save_entry.insert(0, save_folder)

def convert_to_doc():
    pdf_path = pdf_entry.get()
    output_folder = output_entry.get()
    save_location = save_entry.get()

    if pdf_path and output_folder and save_location:
        try:
            images = convert_from_path(pdf_path)
            
            pdf_file = open(pdf_path, 'rb')
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            output_file = f"{output_folder}/converted.docx"
            docx = Document()

            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text = page.extract_text()
                
                paragraph = docx.add_paragraph()
                run = paragraph.add_run(text)
                font = run.font
                font.name = 'Times New Roman'
                font.size = Pt(12)
                
                if page_num < len(images):
                    image_path = f"{output_folder}/image_{page_num + 1}.png"
                    images[page_num].save(image_path)
                    image_text = pytesseract.image_to_string(image_path)
                    
                    paragraph = docx.add_paragraph()
                    run = paragraph.add_run(image_text)
                    font = run.font
                    font.name = 'Times New Roman'
                    font.size = Pt(12)
            
            docx.save(f"{save_location}/converted.docx")
            pdf_file.close()

            result_label.config(text="Conversion successful!", fg="green")
        except Exception as e:
            result_label.config(text=f"Conversion failed: {str(e)}", fg="red")
    else:
        result_label.config(text="Please select PDF file, output folder, and save location.", fg="red")

window = tk.Tk()
window.title("PDF to DOC Converter")

pdf_label = tk.Label(window, text="PDF File:")
pdf_label.pack()

pdf_entry = tk.Entry(window)
pdf_entry.pack()

pdf_button = tk.Button(window, text="Select PDF", command=select_pdf_file)
pdf_button.pack()

output_label = tk.Label(window, text="Output Folder:")
output_label.pack()

output_entry = tk.Entry(window)
output_entry.pack()

output_button = tk.Button(window, text="Select Folder", command=select_output_folder)
output_button.pack()

save_label = tk.Label(window, text="Save Location:")
save_label.pack()

save_entry = tk.Entry(window)
save_entry.pack()

save_button = tk.Button(window, text="Select Save Folder", command=select_save_location)
save_button.pack()

convert_button = tk.Button(window, text="Convert to DOC", command=convert_to_doc)
convert_button.pack()

result_label = tk.Label(window, text="")
result_label.pack()

window.mainloop()

